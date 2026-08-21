from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path


OUT = Path(r"C:\innisfree\Brand-Growth-Strategy\20분_프레젠테이션_페이지별_설계안.docx")

COLORS = {
    "basalt": "253531",
    "forest": "3F5A45",
    "stone": "E8E6D6",
    "rose": "E78093",
    "rose_light": "FBECEF",
    "offwhite": "FAFAF7",
    "ink": "202925",
    "muted": "69736E",
    "line": "D7DDD8",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=COLORS["line"], sz="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Noto Sans KR"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(p, before=0, after=6, line=1.25, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(doc, text, size=10.5, color=COLORS["ink"], bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, before, after, 1.25, align)
    r = p.add_run(text)
    set_run_font(r, size, color, bold, italic)
    return p


def add_labeled(doc, label, text, fill=None, after=5):
    p = doc.add_paragraph()
    style_paragraph(p, 0, after, 1.22)
    if fill:
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)
    r1 = p.add_run(label + " ")
    set_run_font(r1, 10.2, COLORS["forest"], True)
    r2 = p.add_run(text)
    set_run_font(r2, 10.2, COLORS["ink"])
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, 16, COLORS["forest"], True)
    elif level == 2:
        set_run_font(r, 13, COLORS["basalt"], True)
    else:
        set_run_font(r, 11.5, COLORS["forest"], True)
    return p


def add_section_band(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_table_borders(table, COLORS["forest"], "0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["basalt"])
    p = cell.paragraphs[0]
    style_paragraph(p, 0, 0, 1.0)
    r = p.add_run(text)
    set_run_font(r, 11, COLORS["offwhite"], True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_plan(doc, num, title, time, role, p1, p2, p3, emphasis, note=None):
    h = add_heading(doc, f"{num}페이지. {title}", 2)
    add_text(doc, f"발표 시간: {time}  |  페이지 역할: {role}", size=9.5, color=COLORS["muted"], italic=True, after=7)
    add_labeled(doc, "1순위:", p1)
    add_labeled(doc, "2순위:", p2)
    add_labeled(doc, "3순위:", p3)
    add_labeled(doc, "강조:", emphasis, fill=COLORS["rose_light"], after=7)
    if note:
        add_labeled(doc, "주의:", note, fill="FFF7E6", after=9)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, 0, 0, 1.0)
    r = p.add_run("INNISFREE BRAND GROWTH STRATEGY  |  20분 발표 설계안")
    set_run_font(r, 8.5, COLORS["muted"])


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, 0, 0, 1.0)
    r = p.add_run("20분 프레젠테이션 페이지별 설계")
    set_run_font(r, 8.5, COLORS["muted"])


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    set_table_borders(table)
    for idx, head in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, COLORS["stone"])
        p = cell.paragraphs[0]
        style_paragraph(p, 0, 0, 1.05)
        r = p.add_run(head)
        set_run_font(r, 9.3, COLORS["basalt"], True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            style_paragraph(p, 0, 0, 1.12)
            r = p.add_run(str(value))
            set_run_font(r, 9.1, COLORS["ink"])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)
    add_header(sec)
    add_footer(sec)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans KR"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans KR")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans KR")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["forest"], 18, 10),
        ("Heading 2", 13, COLORS["basalt"], 12, 6),
        ("Heading 3", 11.5, COLORS["forest"], 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Noto Sans KR"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans KR")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans KR")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans KR")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    # Cover / proposal centerpiece
    add_text(doc, "INNISFREE BRAND GROWTH STRATEGY", size=11, color=COLORS["rose"], bold=True, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    style_paragraph(p, 0, 8, 1.1, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("20분 프레젠테이션\n페이지별 설계안")
    set_run_font(r, 25, COLORS["basalt"], True)
    add_text(doc, "데이터를 보여주는 발표가 아니라, 데이터를 분석해 도출한 성장 전략을 설득하는 발표", size=12.5, color=COLORS["forest"], after=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_table(doc, ["항목", "설계 방향"], [
        ("발표 목적", "이니스프리 브랜드 인지도 회복 전략의 필요성과 실행안을 설득한다."),
        ("핵심 청중 행동", "헤리티지 재활성화 전략과 3년 실행·측정 원칙을 승인한다."),
        ("발표 중심축", "데이터 → 해석 → 전략적 선택 → 크리에이티브 → 실행 → 승인 요청"),
        ("본편 분량", "26페이지 · 20분"),
    ], [1900, 7460])
    add_text(doc, "작성 기준: 마스터 통합문서·제안서 본문·대시보드 리디자인을 반영한 설득형 발표 설계", size=9.2, color=COLORS["muted"], italic=True, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    add_heading(doc, "1. 발표 설계의 전제", 1)
    add_labeled(doc, "커뮤니케이션 목표:", "청중이 ‘문제는 리브랜딩 자체보다, 바꾼 뒤 소비자에게 닿고 쌓이고 반복되는 실행 구조가 없었던 것’이라고 이해하고, 증명된 헤리티지 전략과 3년 실행 체계를 승인하게 만든다.", fill=COLORS["rose_light"], after=10)
    add_labeled(doc, "발표의 기본 원칙:", "데이터는 결론을 설명하기 위한 근거로만 사용한다. 각 페이지는 지표보다 ‘그래서 어떤 판단을 했는가’를 먼저 보여준다.")
    add_labeled(doc, "핵심 프레임:", "이것은 데이터 리포트가 아니라 전략 제안서다. 본편에서 데이터 원표·방법론·검증 과정은 최소화하고, 전략의 논리와 실행의 구체성을 앞세운다.")
    add_heading(doc, "2. 수정된 전체 시간·페이지 배분", 1)
    add_table(doc, ["구간", "페이지", "시간", "역할"], [
        ("오프닝", "1~3", "1분 50초", "핵심 명제와 승인할 방향을 먼저 제시"),
        ("데이터에서 도출한 진단", "4~10", "4분 20초", "전략이 필요한 이유를 증명"),
        ("전략적 해석", "11~15", "3분 40초", "타깃·기회·포지셔닝을 선택"),
        ("전략을 실제 접점으로 전환", "16~21", "6분", "캠페인·공간·페이지·고객 여정 제안"),
        ("실행·운영·측정", "22~25", "3분 20초", "실행 가능성과 지속 구조를 제시"),
        ("마무리", "26", "50초", "승인 요청과 최종 메시지"),
    ], [3000, 1200, 1400, 3760])
    add_labeled(doc, "비중 조정의 핵심:", "이전안보다 순수 데이터 진단 설명은 압축하고, 전략적 해석·크리에이티브·실행 제안에 더 많은 페이지와 시간을 배정했다.", fill=COLORS["rose_light"], after=10)
    add_heading(doc, "3. 페이지별 상세 설계", 1)
    add_text(doc, "각 페이지는 1순위(화면에 반드시 보여줄 결론), 2순위(결론을 지지하는 근거), 3순위(발표자가 말로 보충하거나 백업으로 이동할 내용), 강조점 순서로 설계한다.", size=10.5, after=10)

    add_section_band(doc, "PART 0  |  오프닝·결론")
    add_page_plan(doc, 1, "이니스프리 브랜드 인지도 회복 제안", "20초", "표지", "이번 발표가 브랜드 성장 전략 제안이라는 점", "제주 헤리티지와 현재 브랜드 상황을 암시하는 최소한의 시각 요소", "팀명·발표자·날짜", "이 발표는 과거를 평가하는 보고서가 아니라 다음 실행을 제안하는 자리다.", "데이터·그래프·목차를 표지에 넣지 않는다.")
    add_page_plan(doc, 2, "리브랜딩은 화제를 만들었지만 브랜드 인지도를 만들지는 못했다", "50초", "핵심 명제", "화제와 브랜드 자산 전환의 차이", "인지 검색 146배, 브랜드 검색 5년 −72.9%, 2026년 1~7월 YoY −31.5%", "리브랜딩 이후 제품 관심은 생겼지만 브랜드 기억으로 쌓이지 않았다는 설명", "문제는 화제 만들기에 실패한 것이 아니라, 화제를 브랜드 자산으로 전환하지 못한 것이다.", "수치를 하나씩 설명하지 말고 발표 전체의 결론을 먼저 심는다.")
    add_page_plan(doc, 3, "오늘 제안하는 것은 ‘제주 복귀’가 아니라 ‘선택의 이유 회복’이다", "40초", "의사결정 예고", "헤리티지 고객 재활성화·증명된 헤리티지·3년 지속 실행", "발표 흐름: 문제 → 전략적 선택 → 실행 제안 → 승인 요청", "청중에게 필요한 결정의 종류", "옛날 이니스프리로 돌아가자는 것이 아니라, 제주를 소비자가 선택하는 근거로 다시 연결하자는 제안이다.", "단순한 목차 페이지 대신 승인할 방향을 먼저 예고한다.")

    add_section_band(doc, "PART 1  |  데이터에서 도출한 진단")
    add_page_plan(doc, 4, "브랜드 관심은 리브랜딩 이후에도 회복되지 않았다", "35초", "문제의 현재 상태", "브랜드 인지도 하락은 일시적 반응이 아니라 지속 추세라는 해석", "브랜드 검색지수 8.503 → 2.302, 5년 −72.9%, 2026년 1~7월 −31.5%", "대시보드의 브랜드 검색 추이 화면", "해결하려는 문제는 디자인 선호도가 아니라 브랜드를 다시 찾게 만드는 구조의 부재다.")
    add_page_plan(doc, 5, "제품에 대한 관심은 브랜드 기억으로 쌓이지 않았다", "40초", "제품-브랜드 괴리", "제품 관심과 브랜드 자산은 별개의 성과라는 해석", "제품군 +118.5% vs 브랜드 −23.9%, 그린티 관심은 단기 스파이크 후 소멸", "제품 검색과 브랜드 검색 추이 비교", "제품을 팔 기회는 만들었지만, 그 제품을 통해 브랜드를 다시 기억하게 만들지는 못했다.", "제품 중심 전략을 그대로 반복하자는 반론을 차단하는 페이지다.")
    add_page_plan(doc, 6, "광고비는 브랜드가 아니라 제품으로 흘렀다", "35초", "자원 연결 실패", "투입과 인지도 성과의 연결 방식이 잘못되었다는 해석", "광고선전비↔브랜드 검색 r=−0.03, 광고선전비↔제품 검색 r=+0.69, 2023년 증가 후 2024년 −20.6% 축소", "광고비 산점도와 제품·브랜드 검색 비교", "광고비를 더 쓰느냐가 아니라, 인지도에 연결되는 예산 구조와 지속 기간이 없었다.", "‘광고비를 늘려서 매출이 줄었다’고 말하지 않는다.")
    add_page_plan(doc, 7, "매장은 판매점이 아니라 브랜드의 무료 미디어였다", "35초", "접점 소멸", "오프라인 접점의 소멸이 브랜드 노출의 소멸이었다는 해석", "매장 400개 → 190개, 매장 수↔브랜드 검색 r=0.98", "매장 수와 브랜드 검색 산점도", "매장을 다시 400개로 늘리자는 것이 아니라, 매장이 하던 브랜드 노출 기능을 대체해야 한다.", "r=0.98을 인과관계로 말하지 않는다.")
    add_page_plan(doc, 8, "제주 자산도, 기능성 전략도 계속 방치할 수 없다", "35초", "자산 감가와 전략 한계", "과거 자산과 현재 전략 모두 지속성이 약해지고 있다는 해석", "2025년 제주 절대 검색이 브랜드보다 빠르게 하락, 제품군도 정점 대비 −37.4%", "제주 vs 브랜드 낙폭 비교와 제품군 정점 이후 하락", "제주를 다시 활용해야 하는 이유는 과거가 아름다워서가 아니라, 사용하지 않으면 사라질 수 있는 자산이기 때문이다.")
    add_page_plan(doc, 9, "AEO 형식은 갖췄지만, 소비자에게 말하는 메시지가 없다", "30초", "디지털 발견 경로", "기술 부재가 아니라 메시지 배치 부재가 문제라는 해석", "AEO 형식 4/4 완비, 타깃 질의어 34개 중 헤리티지 단어 0개", "‘기계에게는 7,595자, 사람에게는 0줄’이라는 대비", "기술은 준비되어 있지만, AI와 소비자가 발견할 수 있는 자리에는 핵심 이야기가 없다.", "‘AEO 대응이 전무하다’ 또는 ‘공식몰에 제주가 없다’고 말하지 않는다.")
    add_page_plan(doc, 10, "다섯 가지 문제는 ‘전달·접점·지속’의 실패로 수렴한다", "50초", "진단 종합과 전략 브리지", "데이터를 하나의 전략적 판단으로 압축", "A~E 문제를 전달·자원 연결·접점·브랜드 축적·발견 경로로 재분류", "문제 A~E와 이후 전략 과제의 매핑", "방향 자체가 완전히 틀린 것이 아니라, 바꾼 뒤 소비자에게 닿고 쌓이고 반복되는 구조가 없었다.", "이 페이지가 데이터 파트와 전략 파트를 연결하는 핵심 브리지다.")

    add_section_band(doc, "PART 2  |  전략적 해석")
    add_page_plan(doc, 11, "가장 먼저 되찾아야 할 사람은 신규 고객이 아니다", "40초", "타깃 선정", "헤리티지를 기억하지만 접점을 잃은 고객을 1순위로 선정", "제주 연상이 완전히 사라지지 않았고, 과거 경험이 있으며, 신규 설득보다 재상기가 빠름", "성분 탐색형은 2순위로 상세페이지에서 대응", "가장 설득하기 쉬운 사람부터 다시 연결해야 한다.")
    add_page_plan(doc, 12, "경쟁사가 비워둔 곳에 이니스프리의 기회가 있다", "35초", "경쟁 기회", "경쟁사보다 더 많이가 아니라, 다르게 말할 수 있다는 판단", "이니스프리·미샤·3CE 등 브랜드 서사 광고 0건, 업계 전체가 제품·할인·유통 중심", "경쟁사 광고 유형 비교", "경쟁사가 하지 않는 방식으로 말하는 것 자체가 차별화가 된다.")
    add_page_plan(doc, 13, "소비자는 제주를 검색하러 오지 않는다. 하지만 제주 때문에 선택할 수 있다", "45초", "소비자 선택 구조", "제주의 역할을 진입 키워드에서 선택 근거로 재정의", "진입: 수분크림·선크림 → 선택: 이니스프리 그린티 → 확신: 제주 원산지와 20년 축적", "소비자 여정 한 줄 구조", "제주는 진입 키워드가 아니라, 일반명사 경쟁에서 최종 선택을 만드는 근거다.", "‘제주를 그린티에 실으면 작동한다’는 표현은 가설로만 다룬다.")
    add_page_plan(doc, 14, "포지셔닝: ‘그때 믿었던 게, 지금도 맞았다’", "50초", "전략 문장", "헤리티지 고객에게 전달할 핵심 가치제안", "기억의 회복, 데이터와 원산지의 근거, 경쟁사와의 차별점", "핵심 키워드 ‘증명된 헤리티지’", "이 문장은 슬로건이 아니라 이후 모든 실행을 평가하는 전략 기준선이다.")
    add_page_plan(doc, 15, "이 전략은 과거 복원이 아니라 헤리티지의 계승이다", "50초", "전략 가드레일", "무엇을 바꾸고 무엇을 유지할지에 대한 기준", "과거 로고·컬러·매장을 전부 되돌리지 않고, 제주를 제품·콘텐츠·공간의 선택 근거로 재배치", "‘과거의 제주’와 ‘증명된 헤리티지’ 비교", "제주로 돌아가는 것이 아니라, 현재의 이니스프리가 제주를 다시 설명할 수 있게 만드는 것이다.")

    add_section_band(doc, "PART 3  |  전략을 실제 접점으로 옮기는 제안")
    add_page_plan(doc, 16, "제안의 중심 아이디어: 증명된 헤리티지를 다시 보이게 하자", "55초", "전략의 실행 언어 전환", "분석 결과를 캠페인·공간·콘텐츠·반복 구조로 번역", "기억: 오랜만이에요 / 근거: 제주 원료와 20년 축적 / 경험: 공병공간 / 반복: 콘텐츠·상세·재방문", "전체 전략 구조를 하나의 흐름으로 표현", "새로운 이야기를 발명하는 것이 아니라, 이미 있는 자산이 소비자에게 보이도록 연결한다.")
    add_page_plan(doc, 17, "캠페인: ‘오랜만이에요’", "1분", "휴면 고객 재활성화", "과거 고객에게 다시 말을 거는 재회형 캠페인", "할인보다 기억과 신뢰 중심, 그때의 성분이 지금도 유효하다는 메시지", "캠페인 무드보드와 실제 카피 예시", "신규 고객을 크게 외치는 캠페인이 아니라, 이미 이니스프리를 알았던 사람에게 다시 말을 거는 캠페인이다.")
    add_page_plan(doc, 18, "공병공간: 이미 있는 자산을 브랜드 접점으로 확장한다", "1분 10초", "오프라인 체험 거점", "매장을 많이 여는 대신 브랜드를 기억하게 만드는 경험을 만든다", "80년 한옥, 공병 23만 개, 그린사이클 23년, 제주 원료와 지속가능성 경험 연결", "기존 공간에서 확장 가능한 체험 요소", "공병공간은 신규 발명이 아니라 이미 있는 자산을 앞으로 꺼내는 실행이다.")
    add_page_plan(doc, 19, "랜딩페이지: 첫 화면에서 브랜드의 이유를 선언한다", "55초", "브랜드 발견 접점", "공식몰을 제품 목록이 아니라 브랜드를 이해하는 입구로 전환", "‘왜 이니스프리인가’ 선언, 제주 원료 20년 타임라인, 수치·FAQ·구조화된 답변", "현재 페이지와 제안 페이지의 Before/After", "공식몰 첫 화면에서 브랜드의 이유를 이해하게 만든다.")
    add_page_plan(doc, 20, "상세페이지: 제품에서 브랜드로 건너가는 다리", "1분", "제품-브랜드 브리지", "제품 검색을 브랜드 재검색과 재방문으로 연결", "일반명사 유입 → 성분·효능 → 제주 원산지·20년 근거 → 브랜드 스토리·공병공간", "상세페이지 구성 목업", "제품을 팔고 끝나는 것이 아니라, 제품을 산 사람이 이니스프리를 다시 기억하게 만든다.")
    add_page_plan(doc, 21, "네 가지 제안은 하나의 고객 여정으로 연결된다", "1분", "통합 성장 구조", "각 제안이 독립 아이디어가 아니라 하나의 고객 경험이라는 점", "일반명사 유입 → 제품 발견 → 제주와 데이터로 선택 → 경험 → 재검색·재방문", "캠페인·상세·랜딩·공병공간 연결 구조", "콘텐츠 하나를 잘 만드는 것이 아니라, 제품 관심이 브랜드 기억으로 이어지는 경로를 설계한다.")

    add_section_band(doc, "PART 4  |  실행·운영·측정")
    add_page_plan(doc, 22, "실행은 세 단계로 시작한다", "55초", "실행 로드맵", "실행 가능성과 우선순위", "0~30일: 채널 감사·플랫폼 제휴 / 30~90일: 캠페인·랜딩·상세·콘텐츠 / 90일 이후: 공병공간 확장", "단계별 선행 조건과 의사결정 포인트", "처음부터 모두 바꾸지 않고, 빠르게 검증할 것과 조직 승인이 필요한 것을 분리한다.")
    add_page_plan(doc, 23, "예산은 ‘더 많이’가 아니라 ‘더 오래, 목적에 맞게’ 써야 한다", "45초", "예산 운영 원칙", "다시 1년 만에 철회되지 않는 운영 구조", "TOFU와 BOFU 예산 분리, 단기 재배치가 아닌 순증, 최소 3년 사전 확약", "2023년 증액 후 1년 만에 축소된 과거 패턴", "이 제안의 예산 처방은 단순 증액이 아니라 지속 가능한 운영 방식이다.", "확정되지 않은 예산 금액은 임의로 넣지 않는다.")
    add_page_plan(doc, 24, "성공은 목표 숫자가 아니라 판정 기준으로 관리한다", "45초", "KPI 원칙", "캠페인 실행 여부가 아니라 자연 추세 대비 개선을 판정", "브랜드 검색지수, 제주 연상 비중, 제품-브랜드 Gap, 직접경쟁 3사 점유율, AEO 실행 지표", "90·180·360일 게이트", "캠페인을 했다는 사실이 아니라, 자연적으로 일어날 변화보다 더 좋아졌는지를 본다.", "AEO는 소비자 인지도 KPI가 아니라 실행 지표로 구분한다.")
    add_page_plan(doc, 25, "대시보드는 제안 이후에도 약속을 지키게 하는 장치다", "55초", "운영 모니터링", "대시보드는 데이터를 보여주는 화면이 아니라 전략 의사결정 장치", "90일 하락폭 축소, 180일 브랜드 연결, 360일 접점·검색 회복", "대시보드 리디자인 화면 한 장", "제안이 끝난 뒤에도 같은 질문을 반복해서 확인하고, 확대·수정·중단 시점을 판단한다.", "전체 화면 투어 대신 핵심 화면 한 장과 판정 로직만 보여준다.")

    add_section_band(doc, "PART 5  |  마무리·승인 요청")
    add_page_plan(doc, 26, "이니스프리는 자산이 부족한 것이 아니라, 자산을 닿게 하는 실행이 부족했다", "50초", "최종 제안과 승인 요청", "헤리티지 고객 재활성화·페이지 개편·공병공간 확장·3년 측정 체계", "오늘 승인받아야 할 세 가지: 전략 방향, 90일 착수, 3년 운영 원칙", "최종 메시지", "오늘 승인받고 싶은 것은 단일 캠페인이 아니라, 자산을 소비자 선택과 브랜드 기억으로 연결하는 3년 실행 체계다.")

    add_heading(doc, "4. 본편에서 줄이거나 백업으로 이동할 내용", 1)
    add_text(doc, "본편은 분석 과정을 전부 증명하는 자리가 아니라, 분석 결과가 왜 이 전략으로 이어지는지 납득시키는 자리다. 다음 내용은 질의응답용 백업 슬라이드로 이동한다.", after=8)
    for item in [
        "PESTEL·5 Forces의 세부 분석",
        "원자료 전체 표와 모든 경쟁 브랜드의 세부 검색 수치",
        "상관계수 산출 과정과 데이터 계열의 상세 차이",
        "제주 연상 정의 변경의 전체 히스토리",
        "AEO 1/4에서 4/4로 바뀐 조사 과정",
        "출처 등급·방법론 세부사항",
        "매출·영업이익을 본편에서 제외한 이유",
    ]:
        add_labeled(doc, "백업 이동:", item, after=3)
    add_heading(doc, "5. 발표 표기 가드", 1)
    add_labeled(doc, "AEO:", "‘대응 전무’가 아니라 ‘형식 4/4 완비, 실제 AI 답변 인용은 별도 검증’으로 표기한다.")
    add_labeled(doc, "제주 연상:", "‘상승했다’가 아니라 ‘브랜드보다 오래 버텼으나 2025년 우위가 역전됐다’고 말한다.")
    add_labeled(doc, "전략 문장:", "‘제주를 그린티에 실으면 작동한다’는 미검증 가설로 표시하고, 확정 전략은 ‘제주는 선택 근거’로 둔다.")
    add_labeled(doc, "경쟁 분모:", "검색 점유율 판정은 직접경쟁 3사(이니스프리·미샤·3CE)로 통일한다.")
    add_labeled(doc, "인과 표현:", "상관은 관찰로만 말하고, 실행 제안은 ‘대체해야 할 기능’과 ‘검증할 가설’로 표현한다.")
    add_heading(doc, "6. 질의응답용 백업 슬라이드 권장 구성", 1)
    add_table(doc, ["백업", "대비할 질문"], [
        ("B1", "이 숫자의 출처와 신뢰도는 무엇인가?"),
        ("B2", "r=0.98이면 매장을 늘리면 되는 것 아닌가?"),
        ("B3", "제주 연상 비중은 정확히 무엇을 뜻하는가?"),
        ("B4", "왜 토니모리를 직접 비교군에서 제외했는가?"),
        ("B5", "자료마다 검색 수치가 다른 이유는 무엇인가?"),
        ("B6", "데이터를 직접 재현할 수 있는가?"),
        ("B7", "왜 매출·영업이익을 본편에서 다루지 않았는가?"),
        ("B8", "이 분석의 가장 큰 한계는 무엇인가?"),
    ], [1800, 7560])
    add_heading(doc, "7. 근거 자료", 1)
    add_text(doc, "본 설계안은 다음 내부 자료의 최신 기준을 반영했다.", after=5)
    for path in [
        "MASTER_통합문서.md",
        "03_기획/제안서_본문.md",
        "03_기획/KPI_판정기준표_초안.md",
        "01_분석결과/핵심발견_발표용.md",
        "01_분석결과/2026-08_공식몰_AI대응_실사.md",
        "05_대시보드/redesign/index.html",
        "05_대시보드/data/dashboard_data.json",
    ]:
        add_text(doc, path, size=9.5, color=COLORS["muted"], after=2)

    doc.core_properties.title = "20분 프레젠테이션 페이지별 설계안"
    doc.core_properties.subject = "이니스프리 브랜드 성장 전략 발표 설계"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "설득형 프레젠테이션 페이지별 우선순위 및 강조점"
    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    build()
